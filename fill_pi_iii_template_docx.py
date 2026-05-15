from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.units import cm
from reportlab.pdfgen.canvas import Canvas

from generate_pi_iii_filled import DBSchemaDiagram, EAPDiagram


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
DOWNLOADS = Path("/Users/thiagolcsalves/Downloads")
TEMPLATE_DOCX = DOWNLOADS / "ModelodeProjetoIntegrador_III_ProfaKadidja.docx"
OUTPUT_DOCX = ROOT / "output" / "docx" / "projeto-integrador-iii-fastinbox.docx"
OUTPUT_PDF = ROOT / "output" / "pdf" / "projeto-integrador-iii-fastinbox-template.pdf"
ASSETS_DIR = ROOT / "tmp" / "pi_iii_assets"


def ensure_dirs() -> None:
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_PDF.parent.mkdir(parents=True, exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)


def find_paragraph(doc: DocumentObject, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Parágrafo não encontrado: {text}")


def find_paragraphs(doc: DocumentObject, text: str) -> list[Paragraph]:
    return [paragraph for paragraph in doc.paragraphs if paragraph.text.strip() == text]


def get_body_style(doc: DocumentObject):
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "Objetivo:" and index + 1 < len(doc.paragraphs):
            return doc.paragraphs[index + 1].style
    return doc.paragraphs[0].style


def clear_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
) -> Paragraph:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if underline is not None:
        run.underline = underline
    return paragraph


def set_runs_text(paragraph: Paragraph, texts: list[str]) -> None:
    if not paragraph.runs:
        paragraph.add_run("")
    while len(paragraph.runs) < len(texts):
        paragraph.add_run("")
    for index, run in enumerate(paragraph.runs):
        run.text = texts[index] if index < len(texts) else ""


def insert_paragraph_after(
    paragraph: Paragraph,
    text: str = "",
    *,
    style=None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    if text:
        run = new_paragraph.add_run(text)
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
    return new_paragraph


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def remove_page_breaks(paragraph: Paragraph) -> None:
    for run in paragraph._p.iter():
        if run.tag == qn("w:r"):
            for child in list(run):
                if child.tag == qn("w:br"):
                    run.remove(child)


def insert_picture_after(
    paragraph: Paragraph,
    image_path: Path,
    *,
    width_cm: float,
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
) -> Paragraph:
    target = insert_paragraph_after(paragraph)
    target.alignment = alignment
    run = target.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return target


def add_table_after(paragraph: Paragraph, rows: int, cols: int, style: str | None = None) -> Table:
    document = paragraph._parent.part.document
    table = document.add_table(rows=rows, cols=cols)
    if style:
        table.style = style
    paragraph._p.addnext(table._tbl)
    return table


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    font = run.font
    font.name = "Times New Roman"
    font.size = Pt(10.5)


def set_table_rows(table: Table, rows: list[list[str]]) -> None:
    for row_idx, values in enumerate(rows):
        if row_idx >= len(table.rows):
            break
        for col_idx, value in enumerate(values):
            if col_idx >= len(table.rows[row_idx].cells):
                break
            set_cell_text(table.rows[row_idx].cells[col_idx], value, bold=row_idx == 0)
    for row_idx in range(len(rows), len(table.rows)):
        for col_idx in range(len(table.rows[row_idx].cells)):
            set_cell_text(table.rows[row_idx].cells[col_idx], "")


def clone_style_from(source: Paragraph, target: Paragraph) -> None:
    target.style = source.style
    target.alignment = source.alignment


def add_bullets_after(anchor: Paragraph, items: list[str], *, style_source: Paragraph) -> Paragraph:
    current = anchor
    for item in items:
        current = insert_paragraph_after(current, f"• {item}", style=style_source.style)
    return current


def render_flowable_pdf_png(flowable, pdf_path: Path, png_path: Path, page_size) -> None:
    width, height = page_size
    canvas = Canvas(str(pdf_path), pagesize=page_size)
    avail_width = width - 3 * cm
    avail_height = height - 3 * cm
    w, h = flowable.wrap(avail_width, avail_height)
    x = (width - w) / 2
    y = (height - h) / 2
    flowable.drawOn(canvas, x, y)
    canvas.save()
    from subprocess import run

    run(
        [
            "pdftoppm",
            "-singlefile",
            "-png",
            str(pdf_path),
            str(png_path.with_suffix("")),
        ],
        check=True,
    )


def build_captioned_contact_sheet(
    items: list[tuple[Path, str]],
    output_path: Path,
    *,
    columns: int,
    thumb_width: int,
    caption_height: int = 60,
    padding: int = 24,
    title: str | None = None,
) -> Path:
    images: list[tuple[Image.Image, str]] = []
    for image_path, caption in items:
        image = Image.open(image_path).convert("RGB")
        ratio = thumb_width / image.width
        thumb_height = int(image.height * ratio)
        image = image.resize((thumb_width, thumb_height))
        images.append((image, caption))

    rows = (len(images) + columns - 1) // columns
    max_height = max(image.height for image, _ in images)
    title_height = 70 if title else 0
    canvas = Image.new(
        "RGB",
        (
            columns * thumb_width + (columns + 1) * padding,
            title_height + rows * (max_height + caption_height) + (rows + 1) * padding,
        ),
        "white",
    )
    draw = ImageDraw.Draw(canvas)
    title_font = ImageFont.load_default()
    font = ImageFont.load_default()

    y_offset = padding
    if title:
        draw.text((padding, y_offset), title, fill="black", font=title_font)
        y_offset += title_height

    for index, (image, caption) in enumerate(images):
        row = index // columns
        col = index % columns
        x = padding + col * (thumb_width + padding)
        y = y_offset + row * (max_height + caption_height)
        canvas.paste(image, (x, y))
        draw.rectangle((x, y, x + thumb_width, y + image.height), outline="black", width=2)
        draw.multiline_text((x, y + image.height + 8), caption, fill="black", font=font, spacing=4)

    canvas.save(output_path)
    return output_path


def build_assets() -> dict[str, Path]:
    eap_pdf = ASSETS_DIR / "eap.pdf"
    eap_png = ASSETS_DIR / "eap.png"
    db_pdf = ASSETS_DIR / "db-schema.pdf"
    db_png = ASSETS_DIR / "db-schema.png"
    render_flowable_pdf_png(EAPDiagram(), eap_pdf, eap_png, landscape(A4))
    render_flowable_pdf_png(DBSchemaDiagram(), db_pdf, db_png, A4)

    prototype_sheet = build_captioned_contact_sheet(
        [
            (DOCS_DIR / "documents" / "evidencias-sprint-2" / "01-home-public.png", "1. Home pública"),
            (DOCS_DIR / "documents" / "evidencias-sprint-2" / "02-login.png", "2. Login por perfil"),
            (DOCS_DIR / "documents" / "evidencias-sprint-2" / "03-nutricionista-dashboard.png", "3. Dashboard do nutricionista"),
            (DOCS_DIR / "documents" / "evidencias-sprint-2" / "10-paciente-landing.png", "4. Landing do paciente"),
            (DOCS_DIR / "documents" / "evidencias-sprint-2" / "12-paciente-pagamento.png", "5. Pagamento"),
            (DOCS_DIR / "documents" / "evidencias-sprint-2" / "05-cozinha-kanban.png", "6. Kanban da cozinha"),
        ],
        ASSETS_DIR / "prototype-sheet.png",
        columns=2,
        thumb_width=430,
        title="Fluxo visual do FastInBox - evidências reais do MVP",
    )

    mobile_sheet = build_captioned_contact_sheet(
        [
            (DOCS_DIR / "documents" / "evidencias-sprint-2" / "10-paciente-landing.png", "Acesso do paciente"),
            (DOCS_DIR / "documents" / "evidencias-sprint-2" / "12-paciente-pagamento.png", "Pagamento"),
            (DOCS_DIR / "documents" / "evidencias-sprint-2" / "11-paciente-status.png", "Status do pedido"),
        ],
        ASSETS_DIR / "mobile-sheet.png",
        columns=1,
        thumb_width=520,
        title="Recorte mobile-first atendido pelo Web Application responsivo",
    )

    web_sheet = build_captioned_contact_sheet(
        [
            (DOCS_DIR / "documents" / "evidencias-sprint-2" / "03-nutricionista-dashboard.png", "Dashboard do nutricionista"),
            (DOCS_DIR / "documents" / "evidencias-sprint-2" / "04-nutricionista-pacientes.png", "Gestão de pacientes"),
            (DOCS_DIR / "documents" / "evidencias-sprint-2" / "05-cozinha-kanban.png", "Operação da cozinha"),
            (DOCS_DIR / "documents" / "evidencias-sprint-2" / "06-admin-dashboard.png", "Dashboard administrativo"),
        ],
        ASSETS_DIR / "web-sheet.png",
        columns=2,
        thumb_width=430,
        title="Principais módulos do Web Application",
    )

    return {
        "eap": eap_png,
        "db": db_png,
        "prototype": prototype_sheet,
        "mobile": mobile_sheet,
        "web": web_sheet,
    }


def populate_cover(doc: DocumentObject) -> None:
    set_runs_text(
        find_paragraph(doc, "PROJETO INTEGRADOR III - TURMA (INFORMAR A TURMA)"),
        ["PROJETO INTEGRADOR III - TURMA (", "2026/1 - NOTURNO", ")", ""],
    )
    set_runs_text(find_paragraph(doc, "NOME DO PROJETO"), ["FASTINBOX"])
    members = find_paragraphs(doc, "RA e Nome")
    cover_members = [
        "RA a informar | Thiago L. C. Alves",
        "RA a informar | Joao Vitor",
        "RA a informar | Gabriel",
        "RA a informar | Integrante 4",
        "RA a informar | Integrante 5",
    ]
    for paragraph, value in zip(members, cover_members):
        set_runs_text(paragraph, [value])
    set_runs_text(find_paragraph(doc, "BRASÍLIA, mmmm de AAAA"), ["BRASÍLIA", ", abril de 2026", ""])
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == "SUMÁRIO" and idx > 0:
            remove_page_breaks(doc.paragraphs[idx - 1])
            break


def populate_glossary(doc: DocumentObject) -> None:
    lead = find_paragraph(doc, "[Insira aqui, em ordem alfabética, a lista de termos, siglas e acrônimos importantes para a compreensão do domínio do projeto/produto]")
    style_source = find_paragraph(doc, "Termo/Sigla")
    set_runs_text(lead, ["[Lista, em ordem alfabética, dos termos, siglas e acrônimos relevantes para a compreensão do domínio do FastInBox.]"])
    lead.runs[0].italic = True
    terms = [
        "API",
        "BMC",
        "EAP",
        "ERS",
        "MVP",
        "PEST",
        "RF",
        "RNF",
        "White label",
    ]
    placeholders = find_paragraphs(doc, "Termo/Sigla")
    for paragraph, value in zip(placeholders, terms):
        set_runs_text(paragraph, [value, ""])
        paragraph.style = style_source.style


def populate_description(doc: DocumentObject) -> None:
    set_paragraph_text(
        find_paragraph(doc, "Informar o objetivo do projeto e não do produto."),
        "Planejar, construir, validar e documentar o MVP acadêmico do FastInBox, aplicando disciplinas de gestão de projetos e de engenharia de software para demonstrar a viabilidade de digitalizar o fluxo de venda, pagamento, produção e acompanhamento de marmitas personalizadas em um ambiente white label.",
    )
    set_paragraph_text(
        find_paragraph(doc, "Descreva o que compõe o projeto e não o produto.  Elencar as disciplinas de gestão bem como as de engenharia de software que serão exercitadas para a construção e entrega do objetivo do projeto."),
        "O projeto reúne os artefatos de negócio, gestão e engenharia necessários à entrega do PI III: Business Model Canvas, Documento de Visão, ERS, Plano de Projeto, Arquitetura UML, Modelo de Dados, protótipos, implementação front-end e baseline de back-end, evidências de sprint e documentação publicada em GitHub Pages. Ao longo do ciclo, são exercitadas disciplinas de Engenharia de Requisitos, Engenharia de Software, Banco de Dados, Programação Web, Interação Humano-Computador, Qualidade de Software e Gestão de Projetos.",
    )


def populate_eap(doc: DocumentObject, assets: dict[str, Path]) -> None:
    paragraph = find_paragraph(doc, "[Utilizar uma ferramenta gráfica para inserir a EAP. Sugestão utilizar o Microsoft Word, Smart Art e Inserir - Hierarquia]")
    set_paragraph_text(paragraph, "A figura a seguir apresenta a EAP do FastInBox, organizada em gestão do projeto, produto e negócio, engenharia e desenvolvimento, validação e evidências e sustentação acadêmica.")
    insert_picture_after(paragraph, assets["eap"], width_cm=16.2)


def populate_problem(doc: DocumentObject) -> None:
    paragraph = find_paragraph(doc, "Fundamente com números e fatos o problema/oportunidade e como ele afeta as pessoas.  Faça de 3 a 5 parágrafos para expor o problema/oportunidade.")
    set_paragraph_text(paragraph, "O processo de venda de marmitas personalizadas em clínicas de nutrição ainda é fragmentado: o nutricionista coleta dados por mensagens, revisa pedidos manualmente, cobra o paciente por links avulsos e repassa a produção para a cozinha sem um fluxo transacional único. Esse cenário aumenta retrabalho, risco de erro operacional e dificuldade de acompanhamento do pedido.")
    p2 = insert_paragraph_after(paragraph, "A documentação do FastInBox mostra que o produto foi desenhado para integrar quatro papéis centrais - nutricionista, paciente, cozinha e administração - em uma única jornada. A oportunidade surge justamente da ausência de uma plataforma que preserve a marca da clínica, mantenha clareza para o paciente e entregue rastreabilidade operacional para produção, pagamento e status.")
    p3 = insert_paragraph_after(p2, "As referências bibliográficas já consolidadas no projeto reforçam o contexto: o avanço da digitalização do comércio e dos meios de pagamento, a ampliação do uso de canais digitais pela população e a pressão por serviços de saúde mais convenientes criam um ambiente favorável para soluções como o FastInBox. O projeto também se apoia na necessidade prática de reduzir improviso operacional em clínicas e cozinhas parceiras.")
    insert_paragraph_after(p3, "Assim, o problema não é apenas tecnológico, mas de coordenação de processo. A oportunidade do FastInBox está em transformar uma operação informal, dispersa e pouco auditável em um fluxo profissional, white label e mensurável, capaz de melhorar a experiência do paciente e elevar a previsibilidade da operação.")


def populate_bmc(doc: DocumentObject) -> None:
    table = doc.tables[0]
    set_cell_text(table.cell(1, 0), "Clínicas de nutrição e nutricionistas parceiros.\nCozinhas de produção e entrega.\nGateway de pagamento.\nHospedagem cloud e serviços de documentação.")
    set_cell_text(table.cell(1, 1), "Desenvolvimento e evolução da plataforma.\nOnboarding de clínicas.\nOperação do checkout.\nGovernança de dados e auditoria.")
    set_cell_text(table.cell(1, 2), "Equipe de produto e engenharia.\nAplicação web em Next.js.\nAPI em NestJS.\nDocumentação pública no GitHub Pages.")
    set_cell_text(table.cell(3, 0), "Plataforma white label que unifica criação do pedido, revisão, pagamento, produção e acompanhamento, reduzindo operação manual e aumentando rastreabilidade para clínica, paciente e cozinha.")
    set_cell_text(table.cell(3, 1), "Onboarding guiado.\nSuporte digital.\nComunicação transacional por jornada.\nPainéis por perfil.")
    set_cell_text(table.cell(3, 2), "Aplicação web responsiva.\nLanding page por código do pedido.\nDashboard da cozinha.\nDashboard administrativo.\nHub documental.")
    set_cell_text(table.cell(5, 0), "Nutricionistas e clínicas.\nPacientes finais.\nCozinhas parceiras.\nAdministração da plataforma.")
    set_cell_text(table.cell(5, 1), "Infraestrutura cloud.\nEvolução do produto.\nSuporte operacional.\nManutenção de integrações e observabilidade.")
    set_cell_text(table.cell(5, 2), "Comissão por pedido.\nAssinatura mensal por clínica.\nPacotes premium de analytics e governança.")


def populate_pest_benefits_public(doc: DocumentObject) -> None:
    set_paragraph_text(
        find_paragraph(doc, "[Cenários de negócio são configurações de futuro que têm probabilidade de acontecimento.  Eles podem apresentar ameaças mas também oportunidades.  Foque nas oportunidades em que estes cenários sejam favoráveis para explorar a solução que será desenvolvida a partir do projeto.  Faça uma análise PEST].  Consultem instituições que fazem estudos e previsão de cenários como: Gartner, Accenture, Mc Kinsey, Deloitte, OCDE, Banco Mundial, BNDES, Movimento  Brasil 6.0, Movimento , Movimento Brasil Competitivo, CLDF, Câmara dos Deputados, Senado Federal, ONGs, Observatórios, Institutos de Pesquisa."),
        "A análise PEST do FastInBox considera o contexto da digitalização de serviços, da expansão dos meios de pagamento instantâneo e da busca por processos mais rastreáveis em saúde e alimentação. O quadro abaixo sintetiza cenários de oportunidade e risco para o projeto.",
    )
    set_table_rows(
        doc.tables[1],
        [
            ["Dimensão", "Pessimista", "Realista", "Otimista"],
            ["Política", "Aumento de exigências regulatórias e tributárias sobre pequenos negócios de alimentação.", "Manutenção do marco regulatório atual com necessidade de adequação gradual.", "Incentivos à digitalização e à saúde digital favorecem adoção de soluções integradas."],
            ["Econômica", "Pressão inflacionária sobre insumos reduz margem de operação.", "Mercado de alimentação saudável segue em crescimento moderado.", "Maior adesão a serviços por assinatura e pedidos recorrentes amplia o mercado endereçável."],
            ["Social", "Baixa maturidade digital de parte das clínicas pode retardar adoção.", "Nutricionistas aderem a ferramentas que reduzam retrabalho com clareza operacional.", "Busca crescente por conveniência e alimentação personalizada acelera a demanda pela solução."],
            ["Tecnológica", "Custos de integração e segurança podem crescer acima do esperado.", "Cloud, gateways e frameworks modernos reduzem barreiras técnicas para o MVP.", "Maturidade de analytics e automação amplia o potencial de escala e diferenciação do produto."],
        ],
    )

    benefits = find_paragraph(doc, "[Descreva de forma clara os benefícios que podem ser alcançados com a solução que será entregue. Lembre da análise PEST e deixe claro o valor para os clientes/usuários da solução proposta.]")
    set_paragraph_text(benefits, "Os benefícios do FastInBox decorrem da centralização do fluxo operacional em uma plataforma única, com impacto direto sobre produtividade, experiência do usuário e governança.")
    benefits_items = [
        "Para o nutricionista: redução do tempo de criação do pedido, padronização comercial e preservação da marca da clínica no modelo white label.",
        "Para o paciente: revisão clara do pedido, pagamento mais simples e acompanhamento transparente do status da produção e da entrega.",
        "Para a cozinha: fila operacional qualificada com pedidos já confirmados e pagos, diminuindo retrabalho e ruído de comunicação.",
        "Para a administração: visão consolidada de usuários, pedidos, pagamentos, eventos e indicadores para tomada de decisão.",
    ]
    last = add_bullets_after(benefits, benefits_items, style_source=benefits)

    public_target = find_paragraph(doc, "[Descreva com mais detalhes o público-alvo identificado no Business Model Canvas]")
    set_paragraph_text(public_target, "O público-alvo do FastInBox é composto por quatro perfis diretamente envolvidos na operação:")
    public_items = [
        "Nutricionistas e clínicas de nutrição que desejam transformar o atendimento em uma operação comercial padronizada, com mais controle e sem perder sua identidade visual.",
        "Pacientes finais que acessam o pedido principalmente pelo celular e precisam de uma jornada simples, confiável e objetiva para confirmar e pagar.",
        "Cozinhas parceiras que operam sob pressão de tempo e dependem de uma fila clara, com status rastreáveis e dados suficientes para execução.",
        "Gestores administrativos da operação, interessados em governança, auditoria, indicadores e previsibilidade para escalar o modelo.",
    ]
    add_bullets_after(public_target, public_items, style_source=public_target)


def populate_schedule(doc: DocumentObject) -> None:
    set_paragraph_text(
        find_paragraph(doc, "[Insira aqui apenas grandes marcos do projeto que tenham relação com as atividades constantes no Trello e de acordo com a EAP]. O cronograma detalhado das atividades e sprints estará na ferramenta TRELLO.COM"),
        "Os marcos abaixo sintetizam o planejamento macro do PI III. O detalhamento operacional por sprint permanece documentado nas páginas de planning e no board de projeto da equipe.",
    )
    set_table_rows(
        doc.tables[2],
        [
            ["Marco", "Data"],
            ["Kickoff e consolidação do escopo, visão e requisitos", "26/02/2026"],
            ["Entrega do MVP navegável da Sprint 1", "10/04/2026"],
            ["Publicação do hub técnico no GitHub Pages", "10/04/2026"],
            ["Planejamento da migração para API e banco relacional", "25/04/2026"],
            ["Entrega da Sprint 2 com reforço de governança e testes", "08/05/2026"],
            ["Fechamento dos testes internos do PI III", "10/07/2026"],
            ["Consolidação do documento principal acadêmico", "24/07/2026"],
            ["Apresentação do Projeto Integrador III", "24/07/2026"],
        ],
    )


def populate_requirements(doc: DocumentObject) -> None:
    set_table_rows(
        doc.tables[3],
        [
            ["Código", "Descrição"],
            ["RF001", "Autenticar usuários por perfil com rotas segregadas para nutricionista, paciente, cozinha e administrador."],
            ["RF002", "Cadastrar e atualizar pacientes vinculados ao nutricionista autenticado."],
            ["RF003", "Criar pedidos com múltiplos itens, observações e código único de acesso para o paciente."],
            ["RF004", "Configurar identidade visual da clínica para operação white label."],
            ["RF005", "Permitir ao paciente acessar o pedido pela landing page usando o código único."],
            ["RF006", "Permitir revisão e confirmação do pedido antes do pagamento."],
            ["RF007", "Registrar pagamento e liberar apenas pedidos válidos para a cozinha."],
            ["RF008", "Exibir acompanhamento de status do pedido ao paciente."],
            ["RF009", "Disponibilizar painel operacional da cozinha com atualização de status."],
            ["RF010", "Registrar trilha de eventos operacionais para auditoria e histórico."],
            ["RF011", "Disponibilizar dashboard administrativo para governança da operação."],
            ["RF012", "Gerenciar cozinhas, parceiros e configurações operacionais."],
            ["RF013", "Calcular comissão por pedido e consolidar visão financeira da operação."],
            ["RF014", "Manter evidências de login, pagamento e alteração de status em auditoria."],
        ],
    )
    set_table_rows(
        doc.tables[4],
        [
            ["Código", "Tipo", "Descrição"],
            ["RNF001", "Performance", "A aplicação deve responder com fluidez nas jornadas críticas e manter tempo adequado de carregamento nas telas principais."],
            ["RNF002", "Usabilidade", "A interface deve ser responsiva, clara e adequada a desktop, tablet e mobile, com navegação orientada por perfil."],
            ["RNF003", "Segurança", "A solução deve evoluir com autenticação segura, tráfego protegido, segregação de acesso e aderência às diretrizes de privacidade."],
            ["RNF004", "Manutenibilidade", "O código deve permanecer modular, tipado e acompanhado de documentação técnica viva no GitHub Pages."],
            ["RNF005", "Confiabilidade", "O fluxo de status do pedido deve manter consistência transacional e rastreabilidade de eventos."],
            ["RNF006", "Disponibilidade", "Os ambientes públicos de demonstração e documentação devem permanecer acessíveis para validação acadêmica."],
            ["RNF007", "Compatibilidade", "A aplicação deve funcionar nos navegadores modernos mais utilizados pelo público-alvo."],
            ["RNF008", "Auditoria", "Eventos sensíveis devem ser registráveis para apoio à governança e análise de incidentes."],
            ["RNF009", "Escalabilidade", "A arquitetura deve permitir a evolução da persistência local para backend e banco relacional sem ruptura conceitual."],
        ],
    )


def populate_prototype(doc: DocumentObject, assets: dict[str, Path]) -> None:
    paragraph = find_paragraph(doc, "[O protótipo visual pode ser feito em diversas ferramentas à sua escolha, como por exemplo o FIGMA ou AXURE.  Insira aqui o link para o mesmo e as imagens em sequência de uso padrão]")
    set_paragraph_text(paragraph, "O protótipo e o MVP visual do FastInBox foram consolidados na documentação oficial do projeto e nas evidências de sprint publicadas em GitHub Pages. Links principais: documentação pública https://fastinbox-repo.github.io/docs/ ; aplicação publicada https://front-chi-flame.vercel.app ; repositório front-end https://github.com/FastInBox-Repo/front .")
    insert_picture_after(paragraph, assets["prototype"], width_cm=11.5)


def populate_mvps(doc: DocumentObject, assets: dict[str, Path]) -> None:
    body_style = get_body_style(doc)
    mobile_heading = find_paragraph(doc, "Aplicativo Móvel")
    mobile_intro = insert_paragraph_after(
        mobile_heading,
        "No recorte atual do PI III, o FastInBox não possui aplicativo móvel nativo. A necessidade móvel é atendida pelo Web Application responsivo, especialmente na jornada do paciente, com foco em acesso por código do pedido, confirmação, pagamento e acompanhamento de status.",
        style=body_style,
    )
    insert_paragraph_after(mobile_intro, "Como evolução para PI IV e ciclos posteriores, o projeto prevê recursos como notificações push, leitura por QR Code e experiência mobile dedicada, caso as validações de campo indiquem prioridade para essa frente.", style=body_style)

    web_heading = find_paragraph(doc, "Web Application")
    web_intro = insert_paragraph_after(
        web_heading,
        "O Web Application é o MVP principal do FastInBox. Ele concentra as jornadas de nutricionista, paciente, cozinha e administração, conectando criação do pedido, revisão, pagamento, produção, auditoria e governança em um único fluxo digital.",
        style=body_style,
    )
    web_bullets = [
        "Autenticação por perfil e rotas segregadas.",
        "Cadastro e gestão de pacientes.",
        "Criação de pedidos com código único.",
        "Landing do paciente para revisão, confirmação e pagamento.",
        "Kanban da cozinha para operação do pedido pago.",
        "Dashboard administrativo com visão consolidada.",
    ]
    last = add_bullets_after(web_intro, web_bullets, style_source=web_intro)
    insert_paragraph_after(last, "As evidências visuais completas dessas jornadas estão publicadas nas seções de sprint review e no documento principal do projeto hospedado em GitHub Pages.", style=body_style)


def populate_data_model(doc: DocumentObject, assets: dict[str, Path]) -> None:
    paragraph = find_paragraph(doc, "[Insira a modelagem da base de dados da aplicação sob a ótica do Backend. Utilizar o MySQL Workbench ou similar]")
    set_paragraph_text(paragraph, "A modelagem abaixo consolida a visão estrutural de backend já descrita no documento Esquema de Banco de Dados do FastInBox, contemplando identidade, clínica, pacientes, pedidos, itens, pagamentos, eventos operacionais, comissões e trilha de auditoria.")
    insert_picture_after(paragraph, assets["db"], width_cm=12.0)


def populate_tests(doc: DocumentObject) -> None:
    internal = find_paragraph(doc, "[Insira o quadro resumo e insira o link para o relatório ou ferramenta de teste]")
    set_paragraph_text(internal, "Os testes internos documentam apenas o que foi efetivamente validado no PI III. Não foram inventados resultados inexistentes; as etapas alfa e beta permanecem explicitamente planejadas para PI IV.")
    table = add_table_after(internal, 5, 3)
    rows = [
        ["Item", "Método", "Resultado"],
        ["Build baseline", "Execução de build dos módulos", "Validado no baseline da sprint"],
        ["Fluxo E2E do MVP", "Navegação manual guiada", "Disponível para demonstração"],
        ["Publicação", "Deploy Vercel + GitHub Pages", "Ambientes públicos ativos"],
        ["Coerência documental", "Revisão cruzada dos artefatos", "Atualizada até abril de 2026"],
    ]
    set_table_rows(table, rows)

    closed = find_paragraph(doc, "[Insira o quadro resumo e insira o link para o relatório ou ferramenta de teste]")
    set_paragraph_text(closed, "Planejado para PI IV. A rodada alfa/fechada será realizada com equipe de teste, convidados acadêmicos e orientação, com roteiro estruturado para avaliar clareza da jornada, entendimento do pedido e leitura dos status operacionais.")
    beta = find_paragraphs(doc, "[Insira o quadro resumo e insira o link para o relatório ou ferramenta de teste]")[-1]
    set_paragraph_text(beta, "Planejado para PI IV. Os testes beta deverão envolver nutricionistas parceiros, pacientes piloto e stakeholders da operação, medindo taxa de confirmação, conversão de pagamento, lead time de produção e principais pontos de atrito da jornada.")


def populate_marketing(doc: DocumentObject) -> None:
    set_paragraph_text(find_paragraph(doc, "Estratégia de Monetização"), "Estratégia de Monetização: comissão por pedido, assinatura mensal por clínica e futura oferta de analytics e governança como camada premium.")
    set_paragraph_text(find_paragraph(doc, "Estratégia de Divulgação do Produto"), "Estratégia de Divulgação do Produto: landing page pública, documentação técnica aberta, demonstrações guiadas e conteúdo digital com foco em produtividade da clínica.")
    set_paragraph_text(find_paragraph(doc, "Estratégia de Aquisição de Clientes"), "Estratégia de Aquisição de Clientes: piloto com clínicas de referência, indicação entre nutricionistas e prova de valor baseada em redução de retrabalho e clareza operacional.")
    set_paragraph_text(find_paragraph(doc, "Estratégia de Formação de Preços"), "Estratégia de Formação de Preços: modelo híbrido com ticket de entrada reduzido, assinatura por faixa de volume e fee variável por pedido após a validação comercial.")
    set_paragraph_text(find_paragraph(doc, "Desdobramento dos 4 Ps para os 4Cs"), "Desdobramento dos 4 Ps para os 4 Cs: Produto -> Cliente; Preço -> Custo; Praça -> Conveniência; Promoção -> Comunicação, sempre com foco na experiência white label e na confiabilidade operacional.")
    set_paragraph_text(find_paragraph(doc, "Vídeo Promocional"), "Vídeo Promocional: planejado para PI IV, apresentando a dor atual do processo manual, o fluxo FastInBox e a proposta de valor para clínicas, nutricionistas e pacientes.")
    set_paragraph_text(find_paragraph(doc, "Vídeo de Instrução de uso do Produto"), "Vídeo de Instrução de uso do Produto: planejado para PI IV em formato tutorial, cobrindo login, criação de pedido, revisão do paciente, pagamento e operação da cozinha.")
    set_paragraph_text(find_paragraph(doc, "Insights de Mercado [Modelagem no Google Looker ou similar]"), "Insights de Mercado: previsto painel analítico em Google Looker Studio ou solução equivalente para acompanhar funil de pedido, taxa de conversão, lead time da cozinha e ticket médio.")


def populate_references(doc: DocumentObject) -> None:
    base = find_paragraph(doc, "[Inserir as referências fazendo a ligação com as seções do trabalho.  Em especial na parte de definição do modelo de negócio, estabelecimento do problema/oportunidade e todo e qualquer conteúdo utilizado e indicado para as atividades técnicas.   Utilizar a ferramenta “Citações” do Google Docs no formato APA]")
    refs = [
        "FASTINBOX. Documento de Visão. GitHub Pages, 2026. Disponível em: https://fastinbox-repo.github.io/docs/documents/documento-visao.html .",
        "FASTINBOX. Business Model Canvas Técnico. GitHub Pages, 2026. Disponível em: https://fastinbox-repo.github.io/docs/documents/business-model-canvas.html .",
        "FASTINBOX. Plano de Projeto Técnico. GitHub Pages, 2026. Disponível em: https://fastinbox-repo.github.io/docs/documents/plano-projeto.html .",
        "FASTINBOX. Especificação de Requisitos de Software. GitHub Pages, 2026. Disponível em: https://fastinbox-repo.github.io/docs/documents/requisitos-software.html .",
        "FASTINBOX. Arquitetura UML da Plataforma. GitHub Pages, 2026. Disponível em: https://fastinbox-repo.github.io/docs/documents/arquitetura-uml.html .",
        "FASTINBOX. Esquema de Banco de Dados. GitHub Pages, 2026. Disponível em: https://fastinbox-repo.github.io/docs/documents/esquema-banco-dados.html .",
        "FASTINBOX. MVP Canvas. Repositório docs, 2026.",
        "BRASIL. Ministério da Saúde. Vigitel Brasil 2023. Brasília, 2024.",
        "IBGE. PNAD Contínua TIC 2023. Rio de Janeiro, 2025.",
        "BANCO CENTRAL DO BRASIL. SPI - Relatório anual 2024. Brasília, 2024.",
        "WORLD BANK. Global Economic Prospects. Washington, DC, 2025.",
    ]
    set_paragraph_text(base, refs[0])
    current = base
    for ref in refs[1:]:
        current = insert_paragraph_after(current, ref, style=base.style)


def populate_appendix(doc: DocumentObject) -> None:
    set_paragraph_text(find_paragraph(doc, "(colocar o link das ferramentas e acesso completo para kadidja.oliveira@ceub.edu.br )"), "(links e status de acesso alinhados ao material publicado em docs; compartilhar acessos privados complementares com kadidja.oliveira@ceub.edu.br, quando aplicável)")
    replacements = {
        "Gestão do Projeto  (Trello)": "Gestão do Projeto (Trello / GitHub Project): board de acompanhamento macro e planning por sprint.",
        "Gestão da Configuração (versionadores como GitLab e/ou Github e Google Drive)": "Gestão da Configuração: GitHub Organization FastInBox-Repo + Google Drive para apoio acadêmico.",
        "Prototipação (Figma)": "Prototipação (Figma): arquivo interno de protótipo; evidências visuais consolidadas no GitHub Pages.",
        "Frontend - Aplicativo": "Frontend - Aplicativo: não se aplica como app nativo no MVP; atendimento mobile via aplicação web responsiva.",
        "Frontend - Aplicação Web": "Frontend - Aplicação Web: Next.js 15 + React 19 + TypeScript + Tailwind CSS.",
        "API Management": "API Management: baseline em NestJS, com evolução planejada para integração completa em PI IV.",
        "Servidor de Aplicação": "Servidor de Aplicação: Vercel para o front-end e guia de deployment documentado para backend.",
        "Sistema Gerenciador de Banco de Dados": "Sistema Gerenciador de Banco de Dados: PostgreSQL planejado para consolidação da persistência.",
        "Automação de testes (QASE.IO)": "Automação de testes: QASE.IO previsto para gestão formal dos casos e evidências; build e validações já realizados no baseline.",
        "Bug Tracking (A escolher, mas pode ser por exemplo o GITLAB)": "Bug Tracking: GitHub Issues e artefatos de sprint review/documentação.",
        "Extração, Transformação e Carga (Ferramenta ou Scripts no Google Drive)": "Extração, Transformação e Carga: scripts e fluxo analítico planejados para PI IV.",
        "Visualização de dados (Google Looker)": "Visualização de dados: Google Looker Studio ou ferramenta equivalente, planejado para insights operacionais.",
        "Editoração de vídeo": "Editoração de vídeo: ferramenta a definir no PI IV para material promocional e tutorial.",
        "Publicação": "Publicação: GitHub Pages para documentação e ambiente público web para demonstração do MVP.",
        "Redes Sociais": "Redes Sociais: canais institucionais previstos para abertura na etapa de marketing do PI IV.",
        "IOT": "IoT: não aplicável ao recorte atual do MVP.",
        "5G": "5G: tecnologia habilitadora indireta, sem dependência arquitetural específica no MVP.",
        "Blockchain": "Blockchain: não aplicável ao recorte atual.",
        "Realidade Aumentada": "Realidade Aumentada: não aplicável ao recorte atual.",
        "Realidade Virtual": "Realidade Virtual: não aplicável ao recorte atual.",
        "Robotic Process Automation": "Robotic Process Automation: fora do escopo do MVP acadêmico.",
        "Machine Learning": "Machine Learning: fora do escopo do MVP atual.",
        "Deep Learning": "Deep Learning: fora do escopo do MVP atual.",
        "Inteligência Artificial": "Inteligência Artificial: não utilizada como componente da solução entregue no PI III.",
        "Business Intelligence": "Business Intelligence: previsto via dashboard analítico em PI IV.",
        "Big Data": "Big Data: fora do escopo do MVP; potencial tema apenas em fase de escala futura.",
    }
    for source, target in replacements.items():
        set_paragraph_text(find_paragraph(doc, source), target)


def set_document_defaults(doc: DocumentObject) -> None:
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)


def main() -> None:
    ensure_dirs()
    assets = build_assets()
    document = Document(str(TEMPLATE_DOCX))
    set_document_defaults(document)
    populate_cover(document)
    populate_glossary(document)
    populate_description(document)
    populate_eap(document, assets)
    populate_problem(document)
    populate_bmc(document)
    populate_pest_benefits_public(document)
    populate_schedule(document)
    populate_requirements(document)
    populate_prototype(document, assets)
    populate_mvps(document, assets)
    populate_data_model(document, assets)
    populate_tests(document)
    populate_marketing(document)
    populate_references(document)
    populate_appendix(document)
    document.save(str(OUTPUT_DOCX))
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
